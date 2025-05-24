
from docx import Document
import os
import uuid

def rewrite_resume(resume_text, suggestions):
    # Append suggestions to original resume content
    suggestion_text = "\n".join(suggestions)
    improved_text = f"{resume_text}\n\n---\nSuggestions to Add:\n{suggestion_text}"

    # Create .docx
    doc = Document()
    doc.add_heading("Improved Resume", level=1)
    for line in improved_text.split("\n"):
        doc.add_paragraph(line)

    # Save with unique filename
    output_dir = "generated_resumes"
    os.makedirs(output_dir, exist_ok=True)
    filename = f"{uuid.uuid4().hex}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)

    return improved_text, filepath
