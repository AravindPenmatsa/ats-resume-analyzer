
from fastapi import APIRouter, UploadFile, File
from fastapi.responses import FileResponse
from app.matcher.keyword_matcher import calculate_ats_score
from app.generator.resume_writer import rewrite_resume
import os
import fitz  # PyMuPDF
from docx import Document

router = APIRouter()

def read_file_content(upload_file: UploadFile):
    try:
        content = upload_file.file.read()
        return content.decode("utf-8")
    except UnicodeDecodeError:
        try:
            return content.decode("windows-1252")
        except:
            return ""

def read_docx(upload_file: UploadFile):
    try:
        doc = Document(upload_file.file)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception:
        return ""

def read_pdf(upload_file: UploadFile):
    try:
        with fitz.open(stream=upload_file.file.read(), filetype="pdf") as doc:
            return "\n".join(page.get_text() for page in doc)
    except Exception:
        return ""

@router.post("/evaluate")
async def evaluate_resume(resume: UploadFile = File(...), jd: UploadFile = File(...)):
    resume_ext = os.path.splitext(resume.filename)[-1].lower()
    jd_ext = os.path.splitext(jd.filename)[-1].lower()

    if resume_ext == ".docx":
        resume_text = read_docx(resume)
    elif resume_ext == ".pdf":
        resume_text = read_pdf(resume)
    else:
        resume_text = read_file_content(resume)

    if jd_ext == ".docx":
        jd_text = read_docx(jd)
    elif jd_ext == ".pdf":
        jd_text = read_pdf(jd)
    else:
        jd_text = read_file_content(jd)

    score_data = calculate_ats_score(resume_text, jd_text)
    improved_resume, filepath = rewrite_resume(resume_text, score_data["missing_keywords"])

    return FileResponse(
        filepath,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=os.path.basename(filepath)
    )

from fastapi import APIRouter, UploadFile, File
from fastapi.responses import JSONResponse
from app.matcher.keyword_matcher import calculate_ats_score
from app.generator.resume_writer import rewrite_resume
import os
import uuid
import fitz  # PyMuPDF
from docx import Document

router = APIRouter()

def read_file_content(upload_file: UploadFile):
    try:
        content = upload_file.file.read()
        return content.decode("utf-8")
    except UnicodeDecodeError:
        try:
            return content.decode("windows-1252")
        except:
            return ""

def read_docx(upload_file: UploadFile):
    try:
        doc = Document(upload_file.file)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception:
        return ""

def read_pdf(upload_file: UploadFile):
    try:
        with fitz.open(stream=upload_file.file.read(), filetype="pdf") as doc:
            return "\n".join(page.get_text() for page in doc)
    except Exception:
        return ""

@router.post("/evaluate-details")
async def evaluate_details(resume: UploadFile = File(...), jd: UploadFile = File(...)):
    resume_ext = os.path.splitext(resume.filename)[-1].lower()
    jd_ext = os.path.splitext(jd.filename)[-1].lower()

    if resume_ext == ".docx":
        resume_text = read_docx(resume)
    elif resume_ext == ".pdf":
        resume_text = read_pdf(resume)
    else:
        resume_text = read_file_content(resume)

    if jd_ext == ".docx":
        jd_text = read_docx(jd)
    elif jd_ext == ".pdf":
        jd_text = read_pdf(jd)
    else:
        jd_text = read_file_content(jd)

    score_data = calculate_ats_score(resume_text, jd_text)
    improved_resume, filepath = rewrite_resume(resume_text, score_data["missing_keywords"])

    download_url = f"/download/{os.path.basename(filepath)}"
    return JSONResponse({
        "ats_score": score_data["score"],
        "missing_keywords": score_data["missing_keywords"],
        "download_url": download_url
    })
