import logging
import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from app.core.config import GENERATED_DIR
from app.services.resume_service import parse_resume_to_structure

logger = logging.getLogger("app")

def clean_text(text: str) -> str:
    """
    Clean text by normalizing whitespace and removing excessive spaces.
    This is especially important for text extracted from .doc files via textutil/antiword.
    """
    if not text:
        return text
    
    # Replace multiple spaces with single space
    text = re.sub(r'  +', ' ', text)
    
    # Remove spaces before punctuation
    text = re.sub(r'\s+([.,;:!?])', r'\1', text)
    
    # Normalize line breaks
    text = re.sub(r'\n\s*\n', '\n', text)
    
    return text.strip()

def add_hyperlink(paragraph, url, text, color="0000EE", underline=True):
    """
    Add a hyperlink to a paragraph.
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r element
    new_run = OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = OxmlElement('w:rPr')

    # Add color if it is given
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'none')
        rPr.append(u)
    else:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink

def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    """
    Set cell margins for a table cell.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    
    for margin, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    
    tcPr.append(tcMar)

def generate_formatted_resume_docx(filename: str, enhanced_text: str, user_info: dict) -> str:
    """
    Generate a professionally formatted resume DOCX with enhanced content.
    """
    try:
        # Parse the enhanced text into structured data
        resume_data = parse_resume_to_structure(enhanced_text, user_info)
        
        document = Document()
        
        # Set margins (0.7 inch)
        sections = document.sections
        for section in sections:
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)
        
        # --- Header Section ---
        name_paragraph = document.add_paragraph()
        name_run = name_paragraph.add_run(resume_data['name'])
        name_run.bold = True
        name_run.font.size = Pt(20)
        name_run.font.name = 'Arial'
        
        if resume_data.get('title'):
            title_paragraph = document.add_paragraph()
            title_run = title_paragraph.add_run(resume_data['title'])
            title_run.font.size = Pt(14)
            title_run.font.name = 'Arial'
            title_paragraph.paragraph_format.space_after = Pt(2)
            
        # Contact Info
        contact_parts = []
        if resume_data.get('location'): contact_parts.append(resume_data['location'])
        if resume_data.get('phone'): contact_parts.append(resume_data['phone'])
        if resume_data.get('email'): contact_parts.append(resume_data['email'])
        
        contact_paragraph = document.add_paragraph()
        contact_run = contact_paragraph.add_run(" | ".join(contact_parts))
        contact_run.font.size = Pt(10)
        contact_run.font.name = 'Arial'
        contact_paragraph.paragraph_format.space_after = Pt(2)
        
        if resume_data.get('linkedin'):
            linkedin_paragraph = document.add_paragraph()
            add_hyperlink(linkedin_paragraph, resume_data['linkedin'], resume_data['linkedin'])
            linkedin_paragraph.paragraph_format.space_after = Pt(12)
        else:
            contact_paragraph.paragraph_format.space_after = Pt(12)

        # --- Sections ---
        for section in resume_data['sections']:
            # Section Title
            section_name = section['name']
            if section_name == 'Projects':
                section_name = 'Professional Experience'
            
            heading = document.add_paragraph()
            heading_run = heading.add_run(section_name.upper())
            heading_run.bold = True
            heading_run.font.size = Pt(14)
            heading_run.font.name = 'Arial'
            heading_run.font.color.rgb = RGBColor(79, 129, 189) # #4F81BD
            
            # Add bottom border to heading
            pPr = heading._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'B0C4DE')
            pBdr.append(bottom)
            pPr.append(pBdr)
            
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(6)

            # Section Content
            if section['type'] == 'professional_experience' or section['type'] == 'projects':
                for entry in section['content']:
                    # Company | Location | Duration line
                    company = entry.get('company', '')
                    location = entry.get('location', '')
                    duration = entry.get('duration', '')
                    
                    # Create a table for the header line to handle alignment
                    table = document.add_table(rows=1, cols=2)
                    table.autofit = False
                    table.allow_autofit = False
                    
                    # Set column widths (approximate)
                    table.columns[0].width = Inches(5.0)
                    table.columns[1].width = Inches(2.0)
                    
                    # Left cell: Company | Location
                    cell_left = table.cell(0, 0)
                    left_p = cell_left.paragraphs[0]
                    
                    company_text = ""
                    if company and location:
                        company_text = f"{company} | {location}"
                    elif company:
                        company_text = company
                    
                    if company_text:
                        run = left_p.add_run(company_text)
                        run.bold = True
                        run.font.size = Pt(11)
                        run.font.name = 'Arial'
                    
                    # Right cell: Duration
                    cell_right = table.cell(0, 1)
                    right_p = cell_right.paragraphs[0]
                    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    if duration:
                        run = right_p.add_run(duration)
                        run.bold = True
                        run.font.size = Pt(10)
                        run.font.name = 'Arial'
                    
                    # Remove table borders and spacing
                    for row in table.rows:
                        for cell in row.cells:
                            set_cell_margins(cell, top=0, bottom=0)
                    
                    # Role
                    if entry.get('role'):
                        role_p = document.add_paragraph()
                        role_run = role_p.add_run(entry['role'])
                        role_run.bold = True
                        role_run.italic = True
                        role_run.font.size = Pt(11)
                        role_run.font.name = 'Arial'
                        role_p.paragraph_format.space_after = Pt(3)
                    
                    # Responsibilities (Bullets)
                    for resp in entry.get('responsibilities', []):
                        p = document.add_paragraph(style='List Bullet')
                        run = p.add_run(resp)
                        run.font.size = Pt(11)
                        run.font.name = 'Arial'
                        p.paragraph_format.space_after = Pt(2)
                    
                    # Environment
                    if entry.get('environment'):
                        env_p = document.add_paragraph()
                        env_label = env_p.add_run("Environment: ")
                        env_label.bold = True
                        env_label.font.size = Pt(9)
                        env_label.font.name = 'Arial'
                        
                        # Clean the environment text
                        env_text_content = clean_text(entry['environment'])
                        env_text = env_p.add_run(env_text_content)
                        env_text.italic = True
                        env_text.font.size = Pt(9)
                        env_text.font.name = 'Arial'
                        
                        env_p.paragraph_format.space_before = Pt(4)
                        env_p.paragraph_format.space_after = Pt(12)
                    else:
                        # Add some space after entry if no environment
                        document.add_paragraph().paragraph_format.space_after = Pt(8)

            elif section['type'] == 'bullets':
                for bullet in section['content']:
                    # Clean the bullet text to remove excessive whitespace
                    cleaned_bullet = clean_text(bullet)
                    if cleaned_bullet:  # Only add non-empty bullets
                        p = document.add_paragraph(style='List Bullet')
                        run = p.add_run(cleaned_bullet)
                        run.font.size = Pt(11)
                        run.font.name = 'Arial'
                        p.paragraph_format.space_after = Pt(2)
                document.add_paragraph().paragraph_format.space_after = Pt(8)

            elif section['type'] == 'plain_text_block':
                # Technical Skills table handling
                if section_name == 'Technical Skills':
                    # Clean the content first
                    content = clean_text(section['content'])
                    
                    # Remove pipe characters that come from .doc extraction
                    content = re.sub(r'\s*\|\s*', ' ', content)
                    
                    # Check if it has colon-separated format
                    if ':' in content:
                        table = document.add_table(rows=0, cols=2)
                        table.autofit = False
                        table.columns[0].width = Inches(2.0)
                        table.columns[1].width = Inches(5.0)
                        
                        lines = content.split('\n')
                        for line in lines:
                            line = line.strip()
                            if line and ':' in line:
                                row_cells = table.add_row().cells
                                parts = line.split(':', 1)
                                category = parts[0].strip()
                                skills = parts[1].strip()
                                
                                # Remove any remaining pipe characters
                                category = category.replace('|', '').strip()
                                skills = skills.replace('|', ',').strip()
                                
                                cat_run = row_cells[0].paragraphs[0].add_run(category)
                                cat_run.bold = True
                                cat_run.font.size = Pt(11)
                                cat_run.font.name = 'Arial'
                                
                                skill_run = row_cells[1].paragraphs[0].add_run(skills)
                                skill_run.font.size = Pt(11)
                                skill_run.font.name = 'Arial'
                        
                        document.add_paragraph().paragraph_format.space_after = Pt(8)
                    else:
                        # Plain text format
                        p = document.add_paragraph(content)
                        p.paragraph_format.space_after = Pt(8)
                else:
                    # Other plain text sections
                    content = clean_text(section['content'])
                    p = document.add_paragraph(content)
                    p.paragraph_format.space_after = Pt(8)

        # Save Document
        output_filename = f"{os.path.splitext(filename)[0]}_formatted.docx"
        output_path = os.path.join(GENERATED_DIR, output_filename)
        document.save(output_path)
        
        logger.info(f"✅ DOCX generated successfully: {output_path}")
        return output_path

    except Exception as e:
        logger.error(f"❌ Error generating DOCX: {e}", exc_info=True)
        raise e
